#!/usr/bin/env python3
"""
guides.py - best-effort reading of guideDoc / suggestedSchema into HINTS for schema generation.

guideDoc (.xlsx/.docx/.json/.txt/.csv/.md) is read as authoritative *reference* field definitions and
folded into the concept-induction prompt as naming guidance (never as sample data). suggestedSchema
(.json) is parsed and its leaf paths exposed so fieldProfiles can target it.

Optional parsers (openpyxl, python-docx) are imported lazily; if absent, that guide is skipped with a
log line rather than failing the whole job.
"""
import json
import logging
import os

log = logging.getLogger("guides")

_MAX_GUIDE_CHARS = 4000


def load_hints(guide_docs, suggested_schema):
    hints = {"guide_block": "", "suggested_schema": None, "suggested_paths": []}

    texts = []
    for path in (guide_docs or []):
        try:
            t = _read_guide(path)
            if t and t.strip():
                texts.append(f"# {os.path.basename(path)}\n{t.strip()}")
                log.info("guide doc read", extra={"file": os.path.basename(path), "count": len(t)})
        except Exception as e:
            log.warning(f"could not read guide doc: {e}", extra={"file": os.path.basename(str(path))})
    hints["guide_block"] = ("\n\n".join(texts))[:_MAX_GUIDE_CHARS]

    if suggested_schema:
        try:
            with open(suggested_schema, encoding="utf-8") as fh:
                obj = json.load(fh)
            hints["suggested_schema"] = obj
            log.info("suggested schema loaded", extra={"file": os.path.basename(str(suggested_schema))})
        except Exception as e:
            log.warning(f"could not read suggested schema: {e}",
                        extra={"file": os.path.basename(str(suggested_schema))})
    return hints


def _read_guide(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".xlsx", ".xlsm"):
        return _read_xlsx(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".json":
        with open(path, encoding="utf-8") as fh:
            return json.dumps(json.load(fh), ensure_ascii=False)
    if ext in (".txt", ".md", ".csv", ".tsv"):
        with open(path, encoding="utf-8", errors="replace") as fh:
            return fh.read()
    log.warning("unsupported guide doc type; skipping", extra={"file": os.path.basename(path)})
    return ""


def _read_xlsx(path):
    try:
        import openpyxl
    except ImportError:
        log.warning("openpyxl not installed; skipping xlsx guide", extra={"file": os.path.basename(path)})
        return ""
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"## sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _read_docx(path):
    try:
        import docx
    except ImportError:
        log.warning("python-docx not installed; skipping docx guide", extra={"file": os.path.basename(path)})
        return ""
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
